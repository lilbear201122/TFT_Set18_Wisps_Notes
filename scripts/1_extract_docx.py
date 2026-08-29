import glob, os, json
import docx

# Purpose: parse the user's Word notes (source docx) into data/wisps_from_docx.json.
# This is the FIRST step of the pipeline -- run 2_match_en_zh.py and 3_merge_final.py after it.
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA = os.path.join(ROOT, 'data')

local_docx = glob.glob(os.path.join(DATA, '*.docx'))
downloads_docx = sorted(glob.glob(r'C:\Users\ggcl3\Downloads\Set18_DataTFT*.docx'), key=os.path.getmtime, reverse=True)
paths = local_docx + downloads_docx
if not paths:
    raise SystemExit('No source docx found in data/ or Downloads. Place the updated Word notes there first.')
path = paths[0]
d = docx.Document(path)
tables = d.tables

def cell_text(cell):
    parts = []
    for p in cell.paragraphs:
        for line in p.text.split('\n'):
            line = line.strip()
            if line:
                parts.append(line)
    return ' | '.join(parts)

def parse_row(vals):
    round_range, name, cat_rarity, cost, effect, notes = vals
    if '·' in cat_rarity:
        category, rarity = [s.strip() for s in cat_rarity.split('·', 1)]
    else:
        category, rarity = cat_rarity.strip(), ''
    effects = [e.strip() for e in effect.split('|') if e.strip()]
    notes_list = [n.strip() for n in notes.split('|') if n.strip() and n.strip() != '—']
    return {
        'round_range': round_range.strip(),
        'name': name.strip(),
        'category': category,
        'rarity': rarity,
        'cost': cost.strip(),
        'effects': effects,
        'notes': notes_list,
    }

def table_rows(t):
    header = [cell_text(c) for c in t.rows[0].cells]
    out = []
    for row in t.rows[1:]:
        vals = [cell_text(c) for c in row.cells]
        out.append(parse_row(vals))
    return out

# table[0] = round price expectation table
t0 = tables[0]
h0 = [cell_text(c) for c in t0.rows[0].cells]
round_price = [dict(zip(h0, [cell_text(c) for c in row.cells])) for row in t0.rows[1:]]

# table[1] = curated "精選" list WITH personal notes (127 rows, covers stage2~partial stage4)
curated = table_rows(tables[1])
note_map = {}
for w in curated:
    key = (w['name'], w['round_range'], w['cost'], w['effects'][0] if w['effects'] else '')
    note_map[key] = w['notes']

# tables[2..6] = full catalog partitioned by earliest stage (2,3,4,5,6) -> 170 unique wisps total
full = []
for idx in range(2, 7):
    full.extend(table_rows(tables[idx]))

for w in full:
    key = (w['name'], w['round_range'], w['cost'], w['effects'][0] if w['effects'] else '')
    if key in note_map:
        w['notes'] = note_map[key]
        w['has_personal_note'] = True
    else:
        w['has_personal_note'] = False

out = {
    'round_price': round_price,
    'wisps': full,
}
out_path = os.path.join(DATA, 'wisps_from_docx.json')
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(out, f, ensure_ascii=False, indent=2)

# round_price.json is consumed directly by 4_build_site.py
with open(os.path.join(DATA, 'round_price.json'), 'w', encoding='utf-8') as f:
    json.dump(round_price, f, ensure_ascii=False, indent=1)

with open(os.path.join(DATA, 'extract_report.txt'), 'w', encoding='utf-8') as f:
    f.write(f'total wisps (full catalog): {len(full)}\n')
    f.write(f'curated (with notes) rows: {len(curated)}\n')
    matched = sum(1 for w in full if w['has_personal_note'])
    f.write(f'matched personal notes: {matched}\n')
    f.write(f'unmatched (no personal note yet): {len(full) - matched}\n')
    cats = sorted(set(w['category'] for w in full))
    rars = sorted(set(w['rarity'] for w in full))
    f.write('categories: ' + ', '.join(cats) + '\n')
    f.write('rarities: ' + ', '.join(rars) + '\n')
    # show a few unmatched ones as sanity check
    f.write('\nsample unmatched:\n')
    for w in full:
        if not w['has_personal_note']:
            f.write(f"  {w['round_range']} | {w['name']} | {w['category']}·{w['rarity']}\n")
